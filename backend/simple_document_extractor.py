"""
Simple and Fast Document Text Extractor
No heavy dependencies - uses lightweight libraries only
"""
import os
from typing import Optional

# Import lightweight PDF and DOCX readers
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠ PyPDF2 not available. Install with: pip install PyPDF2")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠ python-docx not available. Install with: pip install python-docx")


class FastDocumentExtractor:
    """
    Fast and lightweight document text extractor.
    Extracts text from PDF, DOCX, and TXT files without heavy dependencies.
    """
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._extract_txt,
        }
        
        if PDF_AVAILABLE:
            self.supported_formats['.pdf'] = self._extract_pdf
        
        if DOCX_AVAILABLE:
            self.supported_formats['.docx'] = self._extract_docx
        
        print(f"✓ Document extractor initialized")
        print(f"✓ Supported formats: {list(self.supported_formats.keys())}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content or error message
        """
        try:
            # Validate file
            if not os.path.exists(file_path):
                return f"Error: File not found - {file_path}"
            
            # Get file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Check if format is supported
            if file_ext not in self.supported_formats:
                return f"Error: Unsupported file format '{file_ext}'. Supported: {list(self.supported_formats.keys())}"
            
            # Extract text using appropriate method
            extractor_func = self.supported_formats[file_ext]
            text = extractor_func(file_path)
            
            # Return extracted text
            return text.strip() if text else "Error: No text extracted from document"
            
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyPDF2"""
        if not PDF_AVAILABLE:
            return "Error: PDF support not available. Install PyPDF2: pip install PyPDF2"
        
        try:
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx"""
        if not DOCX_AVAILABLE:
            return "Error: DOCX support not available. Install python-docx: pip install python-docx"
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_formats


# Test the extractor
if __name__ == "__main__":
    extractor = FastDocumentExtractor()
    print("\n" + "="*50)
    print("Fast Document Extractor Ready!")
    print("="*50)
